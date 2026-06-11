"""将 Markdown 格式的 AI 总结导出为 Word (.docx)"""
import io
import re
from datetime import datetime

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn


def _set_run_font(run, size_pt: int = 11, bold: bool = False, color: RGBColor | None = None):
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def _add_styled_paragraph(doc: Document, text: str, style: str = "normal"):
    p = doc.add_paragraph()
    run = p.add_run(text)
    if style == "h1":
        p.style = "Heading 1"
        _set_run_font(run, 18, bold=True, color=RGBColor(0x33, 0x33, 0x33))
    elif style == "h2":
        p.style = "Heading 2"
        _set_run_font(run, 14, bold=True, color=RGBColor(0x66, 0x7E, 0xEA))
    elif style == "h3":
        p.style = "Heading 3"
        _set_run_font(run, 12, bold=True, color=RGBColor(0x55, 0x55, 0x55))
    elif style == "bullet":
        p.style = "List Bullet"
        _set_run_font(run, 11)
    elif style == "quote":
        p.paragraph_format.left_indent = Pt(18)
        _set_run_font(run, 10, color=RGBColor(0x88, 0x88, 0x88))
    else:
        _set_run_font(run, 11)
    return p


def _parse_table_row(line: str) -> list[str]:
    cells = [c.strip() for c in line.strip("|").split("|")]
    return cells


def _is_table_separator(line: str) -> bool:
    return bool(re.match(r"^\|?[\s\-:|]+\|?$", line.strip()))


def _clean_inline(text: str) -> str:
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    text = re.sub(r"\*([^*]+)\*", r"\1", text)
    return text.strip()


def markdown_to_docx(content: str, title: str = "AI 智能速览") -> bytes:
    """将 Markdown 文本转为 docx 字节流"""
    doc = Document()

    title_p = doc.add_heading(title, level=0)
    for run in title_p.runs:
        _set_run_font(run, 20, bold=True)

    meta_p = doc.add_paragraph()
    meta_run = meta_p.add_run(f"导出时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    _set_run_font(meta_run, 9, color=RGBColor(0x99, 0x99, 0x99))
    meta_p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    doc.add_paragraph()

    lines = content.replace("\r\n", "\n").split("\n")
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        stripped = line.strip()

        if not stripped or stripped == "---":
            i += 1
            continue

        if stripped.startswith("# "):
            _add_styled_paragraph(doc, _clean_inline(stripped[2:]), "h1")
            i += 1
            continue
        if stripped.startswith("## "):
            _add_styled_paragraph(doc, _clean_inline(stripped[3:]), "h2")
            i += 1
            continue
        if stripped.startswith("### "):
            _add_styled_paragraph(doc, _clean_inline(stripped[4:]), "h3")
            i += 1
            continue

        if stripped.startswith("> "):
            _add_styled_paragraph(doc, _clean_inline(stripped[2:]), "quote")
            i += 1
            continue

        if stripped.startswith("|") and "|" in stripped[1:]:
            table_rows = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                if not _is_table_separator(lines[i]):
                    table_rows.append(_parse_table_row(lines[i]))
                i += 1
            if table_rows:
                col_count = max(len(r) for r in table_rows)
                table = doc.add_table(rows=len(table_rows), cols=col_count)
                table.style = "Table Grid"
                for ri, row in enumerate(table_rows):
                    for ci in range(col_count):
                        cell_text = _clean_inline(row[ci]) if ci < len(row) else ""
                        cell = table.rows[ri].cells[ci]
                        cell.text = cell_text
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                _set_run_font(run, 10, bold=(ri == 0))
                doc.add_paragraph()
            continue

        bullet_match = re.match(r"^[-*]\s+(.*)", stripped)
        if bullet_match:
            _add_styled_paragraph(doc, _clean_inline(bullet_match.group(1)), "bullet")
            i += 1
            continue

        num_match = re.match(r"^\d+\.\s+(.*)", stripped)
        if num_match:
            p = doc.add_paragraph(style="List Number")
            run = p.add_run(_clean_inline(num_match.group(1)))
            _set_run_font(run, 11)
            i += 1
            continue

        _add_styled_paragraph(doc, _clean_inline(stripped))
        i += 1

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def build_export_filename(title: str, file_id: str | None = None) -> str:
    safe = re.sub(r'[\\/:*?"<>|]', "_", title).strip() or "AI智能速览"
    if len(safe) > 50:
        safe = safe[:50]
    suffix = f"_{file_id[:8]}" if file_id else ""
    return f"{safe}{suffix}.docx"
